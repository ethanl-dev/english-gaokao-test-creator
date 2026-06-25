#!/usr/bin/env python3
"""Build a Gaokao-style English test Word document from a JSON spec.

Usage:
    python scripts/build_test_docx.py --spec test_spec.json --out test.docx
"""
import argparse
import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_run_font(run, en_name="Times New Roman", cn_name="宋体", size_pt=12, bold=False):
    run.font.name = en_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), cn_name)
    rFonts.set(qn("w:ascii"), en_name)
    rFonts.set(qn("w:hAnsi"), en_name)


def add_text_with_breaks(paragraph, text, en="Times New Roman", cn="宋体", size=12, bold=False):
    """Add text to a paragraph, converting \n into line breaks."""
    if not text:
        return
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        set_run_font(run, en, cn, size, bold)


def add_paragraph(doc, text, en="Times New Roman", cn="宋体", size=12, bold=False, align=None, first_line_indent=None, left_indent=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    add_text_with_breaks(p, text, en, cn, size, bold)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_text_with_breaks(p, text, "Times New Roman", "黑体", 18, bold=True)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    add_text_with_breaks(p, text, "Times New Roman", "黑体", 12, bold=True)


def add_passage(doc, passage):
    if not passage:
        return
    for block in passage.split("\n\n"):
        if not block.strip():
            continue
        add_paragraph(doc, block.strip(), en="Times New Roman", cn="Times New Roman", size=12, first_line_indent=0)


def render_options_inline(doc, num, options):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_text_with_breaks(p, f"{num}. ", "Times New Roman", "Times New Roman", 12)
    labels = ["A", "B", "C", "D"]
    for i, label in enumerate(labels):
        text = options.get(label, "") if isinstance(options, dict) else options[i] if i < len(options) else ""
        sep = "    "
        add_text_with_breaks(p, f"{label}. {text}{sep}", "Times New Roman", "Times New Roman", 12)


def render_section_content(doc, section):
    """Render a section/subsection object: heading, instructions, passage, questions, etc."""
    if section.get("heading"):
        add_section_heading(doc, section["heading"])
    if section.get("instructions"):
        add_paragraph(doc, section["instructions"], cn="宋体", size=12)

    # 七选五：先显示 7 个选项，再显示带空格的 passage
    if section.get("gap_options"):
        add_paragraph(doc, "选项：", cn="宋体", size=12, bold=True)
        for i, opt in enumerate(section["gap_options"], start=1):
            label = chr(ord("A") + i - 1)
            add_paragraph(doc, f"{label}. {opt}", en="Times New Roman", cn="Times New Roman", size=12, left_indent=0.5)

    if section.get("passage"):
        add_passage(doc, section["passage"])

    if section.get("options"):
        for opt in section["options"]:
            render_options_inline(doc, opt["num"], {k: opt[k] for k in ["A", "B", "C", "D"] if k in opt})

    if section.get("answers"):
        for ans in section["answers"]:
            add_paragraph(doc, f"{ans['num']}. {ans['answer']}", en="Times New Roman", cn="Times New Roman", size=12)

    if section.get("items"):
        for item in section["items"]:
            q_text = f"{item.get('num', '')}. {item.get('question', '')}".strip()
            if q_text:
                add_paragraph(doc, q_text, en="Times New Roman", cn="Times New Roman", size=12)
            opts = item.get("options")
            if opts:
                render_options_inline(doc, "", opts)
            ans = item.get("answer")
            if ans:
                add_paragraph(doc, f"答案：{ans}", cn="宋体", size=12, left_indent=0.5)

    if section.get("writing_prompts"):
        for wp in section["writing_prompts"]:
            num = wp.get("num", "")
            ptype = wp.get("type", "")
            prompt = wp.get("prompt", "")
            word_count = wp.get("word_count", "")
            sample = wp.get("sample", "")
            if num or ptype:
                add_section_heading(doc, f"{num}. {ptype}".strip(". "))
            if prompt:
                add_paragraph(doc, prompt, en="Times New Roman", cn="宋体", size=12)
            if word_count:
                add_paragraph(doc, f"词数要求：{word_count}", cn="宋体", size=12)
            if sample:
                add_paragraph(doc, "参考范文：", cn="宋体", size=12, bold=True)
                add_paragraph(doc, sample, en="Times New Roman", cn="Times New Roman", size=12, left_indent=0.5)

    # Recurse into subsections
    for sub in section.get("subsections", []):
        render_section_content(doc, sub)


def add_basic_info_table(doc, form_data):
    table = doc.add_table(rows=len(form_data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(11.0)

    for i, entry in enumerate(form_data):
        row = table.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]

        label_cell.text = ""
        p = label_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_with_breaks(p, entry.get("label", ""), "Times New Roman", "宋体", 12, bold=True)

        value_cell.text = ""
        p = value_cell.paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        add_text_with_breaks(p, entry.get("value", ""), "Times New Roman", "宋体", 12)


def fill_form_template(template_path, form_data, output_path):
    doc = Document(template_path)
    if not doc.tables:
        print(f"Warning: no tables in template {template_path}; skipping form fill.")
        return
    table = doc.tables[0]

    def normalize(text):
        return text.replace("\n", "").replace(" ", "").replace("\u3000", "").strip()

    label_map = {normalize(entry.get("label", "")): entry.get("value", "") for entry in form_data}

    for row in table.rows:
        if len(row.cells) < 2:
            continue
        cell_text = normalize(row.cells[0].text)
        for label, value in label_map.items():
            if label and (label in cell_text or cell_text in label):
                row.cells[1].text = ""
                p = row.cells[1].paragraphs[0]
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                add_text_with_breaks(p, value, "Times New Roman", "宋体", 12)
                break

    doc.save(output_path)
    print(f"Filled form: {output_path}")


def build_doc(spec, output_path):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Title
    title = spec.get("title", "英语原创题")
    add_title(doc, title)
    if spec.get("subtitle"):
        add_title(doc, spec["subtitle"])

    # Main sections
    for sec in spec.get("sections", []):
        render_section_content(doc, sec)

    # Answer key
    doc.add_page_break()
    add_title(doc, "参考答案")
    for sec in spec.get("answer_key", []):
        add_section_heading(doc, sec.get("section", ""))
        for line in sec.get("lines", []):
            add_paragraph(doc, line, en="Times New Roman", cn="Times New Roman", size=12)

    # Analysis
    doc.add_page_break()
    add_title(doc, "解析")
    for section_name, items in spec.get("analysis", {}).items():
        add_section_heading(doc, section_name)
        for item in items:
            num = item.get("num", "")
            answer = item.get("answer", "")
            explanation = item.get("explanation", "")
            text = f"{num}. {answer}  {explanation}".strip()
            add_paragraph(doc, text, cn="宋体", size=12, first_line_indent=0)

    # Basic info form (embedded)
    doc.add_page_break()
    add_title(doc, "英语原创试题基本信息表")
    form_data = spec.get("basic_info_form", [])
    if form_data:
        add_basic_info_table(doc, form_data)

    doc.save(output_path)
    print(f"Saved test document: {output_path}")

    # Optional separate filled form template
    form_template = spec.get("form_template_path")
    form_output = spec.get("form_output_path")
    if form_template and form_output:
        fill_form_template(form_template, form_data, form_output)


def main():
    parser = argparse.ArgumentParser(description="Build a Gaokao English test docx from JSON spec.")
    parser.add_argument("--spec", required=True, help="Path to test_spec.json")
    parser.add_argument("--out", required=True, help="Output docx path")
    args = parser.parse_args()

    try:
        from docx import Document  # noqa: F401
    except ImportError as e:
        print("Missing dependency. Install with: pip install --user python-docx")
        raise SystemExit(1) from e

    with open(args.spec, "r", encoding="utf-8") as f:
        spec = json.load(f)

    os.makedirs(os.path.dirname(os.path.abspath(args.out)) or ".", exist_ok=True)
    build_doc(spec, args.out)


if __name__ == "__main__":
    main()
